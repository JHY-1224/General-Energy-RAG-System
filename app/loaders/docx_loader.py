from pathlib import Path

from .base_loader import BaseLoader, LoadedDocument


class DocxLoader(BaseLoader):
    def load(self, path: str | Path) -> list[LoadedDocument]:
        source = Path(path)
        try:
            from docx import Document

            document = Document(source)
            content = "\n\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
            parser = "python-docx"
        except ModuleNotFoundError:
            content = f"# {source.stem}\n\nDOCX 已保存，安装 python-docx 后可提取正文与表格。"
            parser = "placeholder"
        return [LoadedDocument(content=content, source=source.name, metadata={"format": "docx", "parser": parser})]
